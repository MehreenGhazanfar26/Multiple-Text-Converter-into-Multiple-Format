import streamlit as st
import PyPDF2
from docx import Document
from fpdf import FPDF
import pandas as pd
from bs4 import BeautifulSoup
import re  # For text sanitization

# Function to clean the text by removing non-XML compatible characters
def sanitize_text(text):
    # Remove NULL bytes and control characters using regex
    return re.sub(r'[\x00-\x1F\x7F-\x9F]', '', text)

# Streamlit App Interface
st.title("Multiple Text Converter into Multiple Format")

# Upload the file in various formats
uploaded_file = st.file_uploader("Choose a file", type=["txt", "pdf", "docx", "csv", "html"])

if uploaded_file is not None:
    file_type = uploaded_file.name.split(".")[-1]  # Get the file extension
    text = ""

    # Process txt files
    if file_type == "txt":
        text = uploaded_file.read().decode("utf-8")
    
    # Process pdf files
    elif file_type == "pdf":
        pdf_reader = PyPDF2.PdfReader(uploaded_file)
        for page_num in range(len(pdf_reader.pages)):
            page = pdf_reader.pages[page_num]
            text += page.extract_text()
    
    # Process docx files
    elif file_type == "docx":
        doc = Document(uploaded_file)
        for para in doc.paragraphs:
            text += para.text

    # Process csv files
    elif file_type == "csv":
        df = pd.read_csv(uploaded_file)
        text = df.to_string()

    # Process html files
    elif file_type == "html":
        html_content = uploaded_file.read()
        soup = BeautifulSoup(html_content, "html.parser")
        text = soup.get_text()

    # Sanitize the text to remove invalid characters
    text = sanitize_text(text)

    # Display the uploaded text on the interface
    st.subheader("Converted Text")
    st.text_area("Here is the converted text:", value=text, height=300)

    # Save the text in different formats
    st.subheader("Save As")
    save_as_txt = st.button("Save as .txt")
    save_as_docx = st.button("Save as .docx")
    save_as_pdf = st.button("Save as .pdf")
    save_as_csv = st.button("Save as .csv")

    # Save as .txt
    if save_as_txt:
        with open("output_text.txt", "w", encoding="utf-8") as file:
            file.write(text)
        st.success("Text saved as 'output_text.txt'")

    # Save as .docx
    if save_as_docx:
        doc = Document()
        doc.add_paragraph(text)
        doc.save("output_text.docx")
        st.success("Text saved as 'output_text.docx'")

    # Save as .pdf
    if save_as_pdf:
        pdf = FPDF()
        pdf.add_page()
        pdf.set_font("Arial", size=12)
        pdf.multi_cell(200, 10, txt=text)
        pdf.output("output_text.pdf")
        st.success("Text saved as 'output_text.pdf'")

    # Save as .csv
    if save_as_csv:
        df = pd.DataFrame([text], columns=["Text"])
        df.to_csv("output_text.csv", index=False)
        st.success("Text saved as 'output_text.csv'")
